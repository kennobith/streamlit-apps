
import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pdfminer.high_level import extract_text
import re
from io import BytesIO

def obtener_expdte(filename):
    expdte_pdf = filename.split(' - ')[1]
    expdte = expdte_pdf.split('.')[0]
    return expdte

def obtener_reso(filename):
    reso = filename.split(' - ')[0]
    return reso

def extraer_reso_y_expediente(texto):
    texto = texto.strip()

    # Buscar dónde empieza el nroDeExpediente
    match = re.search(r'\bEX-\d{4}-', texto)
    if not match:
        return None, None

    inicio_expte = match.start()

    # Cortamos sin limpiar a ciegas
    nro_reso = texto[:inicio_expte].rstrip(" -")  # solo eliminamos guiones o espacios al final del reso
    nro_expte = texto[inicio_expte:].lstrip(" -") # solo eliminamos guiones o espacios al principio del expte
    nro_expte = nro_expte.split('.pdf')[0]

    return nro_reso, nro_expte

def obtener_nombres_y_legajos(file):
    trabajadores = {} #key = legajo, clave = nombre completo
    texto = extract_text(file)
    texto = texto.replace('\n', ' ') 
    texto = re.sub(r'\s+', ' ', texto)
    patron = r"(?:Dr\.|Dra\.|Lic\.|Ing\.|Sr\.|Sra\.|Prof\.|Mg\.)?\s*" \
         r"([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+" \
         r"(?:\s+(?:[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+|[a-záéíóúñ]+|[A-ZÁÉÍÓÚÑ]+))*)\s*" \
         r"\(Legajo\s+(?:N°|Nº|No|N\.°|N\.º|\s|Num\.?)\s+" \
         r"(\d{1,3}(?:\.\d{3})*|\d+)\)"

    coincidencias = re.findall(patron, texto)

    for nombre, legajo in coincidencias:
        legajo = legajo.replace('.', '')
        trabajadores[legajo] = nombre
    
    return trabajadores

def obtener_datos(directorio):
    datos_expdtes = {}

    for file in st_archivos:
        st.write(f"Vamos con {file.name}")
        #reso = obtener_reso(file.name)
        #expdte = obtener_expdte(file.name)
        reso,expdte = extraer_reso_y_expediente(file.name)
        trabajadores = obtener_nombres_y_legajos(file)
        datos_expdtes[file.name] = {
                                        "expdte": expdte,
                                        "reso": reso,
                                        "trabajadores": trabajadores
                                    }
    return datos_expdtes

def armar_hoja(documento,expediente,resolucion,trabajadores):
     # Encabezado del expediente centrado y en negrita
    parrafo_exp = documento.add_paragraph()
    run_exp = parrafo_exp.add_run(expediente)
    run_exp.bold = True
    run_exp.underline = True
    run_exp.font.size = Pt(16)
    parrafo_exp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Encabezado de la resolución centrado, subrayado y en negrita
    parrafo_res = documento.add_paragraph()
    run_res = parrafo_res.add_run(resolucion)
    run_res.bold = True
    run_res.underline = True
    run_res.font.size = Pt(16)
    parrafo_res.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Tabla
    tabla = documento.add_table(rows=1, cols=3)
    tabla.style = 'Table Grid'  # Bordes visibles

    encabezado = tabla.rows[0].cells
    encabezado[0].text = "LEGAJO"
    encabezado[1].text = "NOMBRE"
    encabezado[2].text = "FIRMA Y FECHA"

    for legajo, nombre in trabajadores.items():
        fila = tabla.add_row().cells
        fila[0].text = str(legajo)
        fila[1].text = nombre
        fila[2].text = ""  # Vacío para firma y fecha

    documento.add_page_break()

def armar_documento(dict,st_files):
    documento = Document()

    style = documento.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    for file in st_files:
        expdte = dict[file.name]["expdte"]
        reso = dict[file.name]["reso"]
        trabajadores = dict[file.name]["trabajadores"]
        armar_hoja(documento,expdte,reso,trabajadores)
    
    return documento

## Streamlit APP ## 

st.title('NotificApp 🧢')
st_archivos = st.file_uploader("Clickeá donde dice 'Browse files' y subí los archivos", accept_multiple_files=True)
st.warning('Recordá que el formato tiene que ser \'NroDeResolucion - NroDeExpediente\'', icon="⚠️")

if st_archivos:
    st.success(f"Subiste {len(st_archivos)} expedientes(s)")
    
    if st.button("Procesar y convertir a DOCX"):
        dict_file_datos = obtener_datos(st_archivos)
        
        documento = armar_documento(dict_file_datos,st_archivos)

        buffer = BytesIO()
        documento.save(buffer)
        buffer.seek(0)

        st.info('Recordá revisar el documento, pues puede contener errores de tipeo')
        st.download_button(
            label="Descargar notificaciones",
            data=buffer,
            file_name="LISTO PARA NOTIFICAR.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            icon=":material/download:",

        )
