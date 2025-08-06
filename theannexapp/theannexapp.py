import openpyxl
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.shared import Mm
import streamlit as st
from io import BytesIO


def armar_anexo(documento,planilla):
    # Encabezado del anexo centrado y en negrita
    parrafo_exp = documento.add_paragraph()
    nombre_anexo = planilla.name.split('.xlsx')[0]
    run_exp = parrafo_exp.add_run(nombre_anexo)
    run_exp.bold = True
    run_exp.underline = True
    run_exp.font.size = Pt(16)
    parrafo_exp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Tabla
    tabla = documento.add_table(rows=1, cols=8)
    tabla.style = 'Table Grid'  # Bordes visibles

    encabezado = tabla.rows[0].cells
    encabezado[0].text = "LEGAJO"
    encabezado[1].text = "APELLIDO Y NOMBRE"
    encabezado[2].text = "FUNCIÓN"
    encabezado[3].text = "Cat."
    encabezado[4].text = "BONIF."
    encabezado[5].text = "INGRESO"
    encabezado[6].text = "EGRESO"
    encabezado[7].text = "NOTIFICACION FIRMA Y FECHA"

    wb = openpyxl.load_workbook(planilla,read_only = True)
    ws = wb.worksheets[0]

    for row in ws.iter_rows(min_row = 2, max_row = ws.max_row, min_col = 3, max_col = 9):
        fila = tabla.add_row().cells
        i = 0
        for cell in row:
            if cell.value is not None:
                fila[i].text = str(cell.value)
            else:
                fila[i].text = ""
            i += 1
        fila[7].text = "" # espacio para firmar

    documento.add_page_break()

def armar_anexos(planillas):
    documento = Document()

    style = documento.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    section = documento.sections[0]
    section.page_height = Mm(210)
    section.page_width = Mm(297)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    section.orientation = WD_ORIENT.LANDSCAPE

    for planilla in planillas:
        armar_anexo(documento,planilla)

    return documento

## Streamlit APP ## 

st.title('The Annex App📎')
st_archivos = st.file_uploader("Clickeá donde dice 'Browse files' y subí los archivos", accept_multiple_files=True)

if st_archivos:
    st.success(f"Subiste {len(st_archivos)} planilla(s)")
    titulo = st.text_input("Escribí el nombre del archivo y presioná Enter", "Anexo Subsecretaría ABC")

    if st.button("Procesar y armar anexos"):
        documento = armar_anexos(st_archivos)

        buffer = BytesIO()
        documento.save(buffer)
        buffer.seek(0)

        st.info('Recordá revisar el documento')
        st.download_button(
            label="Descargar anexos",
            data=buffer,
            file_name= titulo.strip() + ".docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            icon=":material/download:",
        )
